"""
PDF to Word Converter - Converts PDF files to editable Word documents (.docx)
Author: Your Name
Description: Extract text from PDF and save as a formatted Word document
Usage: python pdf_to_word.py --input file.pdf --output result.docx
         python pdf_to_word.py --folder ./pdfs  (convert entire folder)

Requirements:
    pip install pymupdf python-docx
"""

import argparse
import os
import sys
from datetime import datetime

# Dependency check
try:
    import fitz  # PyMuPDF
except ImportError:
    print("❌ Missing library: PyMuPDF")
    print("   Install it with: pip install pymupdf")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Missing library: python-docx")
    print("   Install it with: pip install python-docx")
    sys.exit(1)


def convert_pdf_to_word(pdf_path, output_path=None):
    """
    Convert a single PDF file to a Word document.

    Args:
        pdf_path: Path to the input PDF file
        output_path: Path for the output .docx file (optional)

    Returns:
        Path to the created .docx file, or None on failure
    """
    if not os.path.exists(pdf_path):
        print(f"❌ File not found: {pdf_path}")
        return None

    if output_path is None:
        base = os.path.splitext(pdf_path)[0]
        output_path = base + ".docx"

    print(f"\n📄 Converting: {os.path.basename(pdf_path)}")

    try:
        pdf_doc = fitz.open(pdf_path)
        word_doc = Document()

        # Document title style
        title = word_doc.add_heading(
            f"Converted from: {os.path.basename(pdf_path)}", level=1
        )
        title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        word_doc.add_paragraph(
            f"Converted on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ).runs[0].font.size = Pt(9)

        word_doc.add_paragraph("─" * 60)

        total_pages = len(pdf_doc)
        print(f"   Pages found: {total_pages}")

        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            text = page.get_text("text")

            # Page header
            page_heading = word_doc.add_heading(f"Page {page_num + 1}", level=2)
            page_heading.runs[0].font.color.rgb = RGBColor(0x27, 0x6F, 0xBF)

            if text.strip():
                # Split into paragraphs and add them
                paragraphs = text.split("\n\n")
                for para_text in paragraphs:
                    clean = para_text.strip()
                    if clean:
                        para = word_doc.add_paragraph(clean)
                        para.runs[0].font.size = Pt(11) if para.runs else None
            else:
                word_doc.add_paragraph(
                    "[This page contains images or non-text content]"
                ).runs[0].italic = True

            # Page separator
            if page_num < total_pages - 1:
                word_doc.add_paragraph("")

        pdf_doc.close()
        word_doc.save(output_path)

        size_kb = os.path.getsize(output_path) / 1024
        print(f"   ✅ Saved: {output_path} ({size_kb:.1f} KB)")
        return output_path

    except Exception as e:
        print(f"   ❌ Error converting {pdf_path}: {e}")
        return None


def convert_folder(folder_path, output_folder=None):
    """Convert all PDFs in a folder."""
    if not os.path.exists(folder_path):
        print(f"❌ Folder not found: {folder_path}")
        return

    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".pdf")]

    if not pdf_files:
        print("⚠️  No PDF files found in the folder.")
        return

    if output_folder:
        os.makedirs(output_folder, exist_ok=True)

    print(f"\n📂 Found {len(pdf_files)} PDF file(s) in '{folder_path}'")

    success = 0
    for filename in pdf_files:
        pdf_path = os.path.join(folder_path, filename)
        if output_folder:
            out_name = os.path.splitext(filename)[0] + ".docx"
            output_path = os.path.join(output_folder, out_name)
        else:
            output_path = None

        result = convert_pdf_to_word(pdf_path, output_path)
        if result:
            success += 1

    print(f"\n📊 Done! {success}/{len(pdf_files)} file(s) converted successfully.")


def main():
    parser = argparse.ArgumentParser(
        description="📄 PDF to Word Converter — Convert PDF files to editable .docx"
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input", type=str, help="Path to a single PDF file")
    group.add_argument("--folder", type=str, help="Path to a folder of PDF files")

    parser.add_argument("--output", type=str, help="Output file or folder path")

    args = parser.parse_args()

    if args.input:
        convert_pdf_to_word(args.input, args.output)
    elif args.folder:
        convert_folder(args.folder, args.output)


if __name__ == "__main__":
    main()
